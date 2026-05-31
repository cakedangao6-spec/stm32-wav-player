from __future__ import annotations

import csv
import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parent
FIGURES = ROOT / "figures"
MANUAL_MD = ROOT / "STM32_MP3硬件设计学习手册.md"
OUTPUT_DOCX = ROOT / "STM32_MP3硬件设计学习手册.docx"
CONNECTION_CSV = ROOT / "STM32_MP3原理图连接表.csv"
BOM_CSV = ROOT / "STM32_MP3_BOM.csv"
CONNECTION_XLSX = ROOT / "STM32_MP3原理图连接表.xlsx"
BOM_XLSX = ROOT / "STM32_MP3_BOM.xlsx"

DOC_SKILL_DIR = Path(r"D:\codex-home\plugins\cache\openai-primary-runtime\documents\26.513.11550\skills\documents")
sys.path.append(str(DOC_SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


FONT_REGULAR = r"C:\Windows\Fonts\msyh.ttc"
FONT_BOLD = r"C:\Windows\Fonts\msyhbd.ttc"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size=size)


def draw_centered_text(draw, box, text, text_font, fill="#16324F", spacing=8):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = [draw.textbbox((0, 0), line, font=text_font)[3] for line in lines]
    total_h = sum(heights) + spacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=text_font)
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=text_font, fill=fill)
        y += h + spacing


def rounded_box(draw, box, fill, outline="#16324F", radius=18, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#375A7F", width=5):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        pts = [(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)]
    else:
        pts = [(x2, y2), (x2 + 18, y2 - 10), (x2 + 18, y2 + 10)]
    draw.polygon(pts, fill=fill)


def save_system_block():
    img = Image.new("RGB", (1600, 900), "#F7F8FA")
    d = ImageDraw.Draw(img)
    title = font(38, True)
    body = font(27)
    small = font(22)
    d.text((60, 40), "系统框图", font=title, fill="#16324F")

    boxes = {
        "power": (80, 160, 360, 300),
        "mcu": (520, 160, 820, 300),
        "flash": (980, 120, 1260, 250),
        "oled": (980, 290, 1260, 420),
        "audio": (500, 480, 840, 650),
        "speaker": (1040, 500, 1380, 630),
    }
    rounded_box(d, boxes["power"], "#DDEAF7")
    rounded_box(d, boxes["mcu"], "#D9F0E1")
    rounded_box(d, boxes["flash"], "#FFF0D9")
    rounded_box(d, boxes["oled"], "#FFF0D9")
    rounded_box(d, boxes["audio"], "#F9E0E0")
    rounded_box(d, boxes["speaker"], "#FDECC8")

    draw_centered_text(d, boxes["power"], "5V 输入\n3.3V 稳压", body)
    draw_centered_text(d, boxes["mcu"], "STM32F103C8T6\n控制核心", body)
    draw_centered_text(d, boxes["flash"], "W25Q64\nSPI Flash", body)
    draw_centered_text(d, boxes["oled"], "OLED\nI2C 显示", body)
    draw_centered_text(d, boxes["audio"], "PB0 PWM\n两级 RC + PAM8403", body)
    draw_centered_text(d, boxes["speaker"], "8Ω 小喇叭", body)

    arrow(d, (360, 230), (520, 230))
    arrow(d, (820, 210), (980, 185))
    arrow(d, (820, 250), (980, 355))
    arrow(d, (670, 300), (670, 480))
    arrow(d, (840, 565), (1040, 565))

    d.text((394, 190), "3V3_D", font=small, fill="#375A7F")
    d.text((855, 150), "SPI1", font=small, fill="#375A7F")
    d.text((855, 332), "I2C", font=small, fill="#375A7F")
    d.text((705, 382), "PB0 PWM", font=small, fill="#375A7F")
    d.text((892, 530), "功率音频", font=small, fill="#375A7F")

    img.save(FIGURES / "系统框图.png")


def save_module_map():
    img = Image.new("RGB", (1700, 1100), "#F7F8FA")
    d = ImageDraw.Draw(img)
    title = font(38, True)
    body = font(25)
    small = font(21)
    d.text((60, 40), "模块连接图", font=title, fill="#16324F")

    boxes = {
        "mcu": (620, 380, 1080, 720),
        "flash": (80, 160, 420, 360),
        "oled": (80, 740, 420, 940),
        "uart": (1260, 120, 1580, 300),
        "swd": (1260, 350, 1580, 530),
        "audio": (1260, 640, 1580, 860),
        "speaker": (1260, 900, 1580, 1030),
    }
    for key in boxes:
        fill = {
            "mcu": "#D9F0E1",
            "flash": "#FFF0D9",
            "oled": "#FFF0D9",
            "uart": "#DDEAF7",
            "swd": "#DDEAF7",
            "audio": "#F9E0E0",
            "speaker": "#FDECC8",
        }[key]
        rounded_box(d, boxes[key], fill)

    draw_centered_text(d, boxes["mcu"], "STM32F103C8T6\n\nPA4 CS\nPA5 SCK\nPA6 MISO\nPA7 MOSI\nPB8 SCL\nPB9 SDA\nPA9 TX\nPA10 RX\nPB0 PWM", body)
    draw_centered_text(d, boxes["flash"], "W25Q64\nPA4~PA7", body)
    draw_centered_text(d, boxes["oled"], "SSD1306 OLED\nPB8 / PB9", body)
    draw_centered_text(d, boxes["uart"], "USART1\n文件传输 / 日志", body)
    draw_centered_text(d, boxes["swd"], "SWD\n下载 / 调试", body)
    draw_centered_text(d, boxes["audio"], "两级 RC\n音量\nPAM8403", body)
    draw_centered_text(d, boxes["speaker"], "喇叭接口\nOUT+ / OUT-", body)

    arrow(d, (420, 260), (620, 470))
    arrow(d, (420, 840), (620, 650))
    arrow(d, (1080, 460), (1260, 210))
    arrow(d, (1080, 550), (1260, 440))
    arrow(d, (1080, 640), (1260, 740))
    arrow(d, (1420, 860), (1420, 900))

    d.text((465, 278), "SPI1", font=small, fill="#375A7F")
    d.text((470, 777), "I2C", font=small, fill="#375A7F")
    d.text((1110, 300), "PA9 / PA10", font=small, fill="#375A7F")
    d.text((1110, 474), "PA13 / PA14", font=small, fill="#375A7F")
    d.text((1110, 700), "PB0", font=small, fill="#375A7F")

    img.save(FIGURES / "模块连接图.png")


def save_schematic_sketch():
    img = Image.new("RGB", (1900, 1200), "#F7F8FA")
    d = ImageDraw.Draw(img)
    title = font(38, True)
    body = font(24)
    small = font(19)
    d.text((60, 40), "原理图草图", font=title, fill="#16324F")

    blocks = [
        ((60, 150, 430, 350), "#DDEAF7", "电源\nJ1 / F1 / D1\nME6211 3.3V"),
        ((500, 120, 1020, 430), "#D9F0E1", "STM32 最小系统\n8MHz 晶振\n复位 / BOOT / SWD\n去耦 / VDDA"),
        ((1090, 130, 1450, 310), "#FFF0D9", "W25Q64\nPA4~PA7\nCS/WP/HOLD 上拉"),
        ((1090, 350, 1450, 530), "#FFF0D9", "OLED\nPB8 / PB9\n4.7k 上拉"),
        ((500, 560, 980, 820), "#F9E0E0", "音频前端\nPB0 -> 2.2k/10nF\n两级 RC -> RV1"),
        ((1080, 600, 1450, 850), "#F9E0E0", "PAM8403\n5V_AUDIO\n输入耦合\nSHDN 上拉"),
        ((1540, 650, 1840, 800), "#FDECC8", "喇叭\nOUT+ / OUT-"),
        ((500, 900, 980, 1080), "#E9EEF4", "调试点\n3V3 / GND\nPWM / FILT / NRST"),
    ]
    for box, fill, label in blocks:
        rounded_box(d, box, fill)
        draw_centered_text(d, box, label, body)

    arrow(d, (430, 250), (500, 250))
    arrow(d, (1020, 220), (1090, 220))
    arrow(d, (1020, 340), (1090, 440))
    arrow(d, (760, 430), (760, 560))
    arrow(d, (980, 690), (1080, 690))
    arrow(d, (1450, 725), (1540, 725))
    arrow(d, (740, 820), (740, 900))

    d.text((445, 210), "3V3_D", font=small, fill="#375A7F")
    d.text((1038, 180), "SPI1", font=small, fill="#375A7F")
    d.text((1038, 380), "I2C", font=small, fill="#375A7F")
    d.text((785, 480), "PB0 PWM", font=small, fill="#375A7F")
    d.text((1008, 646), "AUDIO_VOL", font=small, fill="#375A7F")
    d.text((1465, 680), "桥接输出", font=small, fill="#375A7F")

    img.save(FIGURES / "原理图草图.png")


def save_pcb_sketch():
    img = Image.new("RGB", (1700, 1000), "#F7F8FA")
    d = ImageDraw.Draw(img)
    title = font(38, True)
    body = font(25)
    small = font(20)
    d.text((60, 40), "PCB 布局草图", font=title, fill="#16324F")

    board = (80, 130, 1620, 900)
    d.rounded_rectangle(board, radius=28, fill="#FFFFFF", outline="#16324F", width=5)

    zones = [
        ((120, 180, 420, 360), "#DDEAF7", "电源区\nJ1 / LDO"),
        ((520, 220, 980, 620), "#D9F0E1", "STM32 区\n晶振贴近\n去耦贴脚"),
        ((1040, 210, 1300, 400), "#FFF0D9", "W25Q64\n靠近 SPI"),
        ((120, 660, 410, 840), "#FFF0D9", "OLED 接口\n靠板边"),
        ((1030, 520, 1450, 820), "#F9E0E0", "音频区\nRC / RV1 / PAM8403"),
        ((1480, 590, 1580, 760), "#FDECC8", "喇叭\n端子"),
    ]
    for box, fill, label in zones:
        rounded_box(d, box, fill)
        draw_centered_text(d, box, label, body)

    d.text((536, 646), "连续完整 GND 面，不随意切地", font=small, fill="#375A7F")
    d.text((1005, 470), "低电平音频线短且远离 SPI / 晶振", font=small, fill="#375A7F")
    d.text((519, 180), "SWD / UART 靠板边", font=small, fill="#375A7F")

    arrow(d, (420, 270), (520, 270))
    arrow(d, (980, 300), (1040, 300))
    arrow(d, (820, 620), (1030, 620))
    arrow(d, (1450, 675), (1480, 675))

    img.save(FIGURES / "PCB布局草图.png")


def generate_figures():
    FIGURES.mkdir(parents=True, exist_ok=True)
    save_system_block()
    save_module_map()
    save_schematic_sketch()
    save_pcb_sketch()


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10.0):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_run_font(
    run,
    size: float = 11,
    bold: bool | None = None,
    color: RGBColor | None = None,
    east_asia: str = "宋体",
    latin: str = "Calibri",
):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_inline_runs(paragraph, text: str):
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, east_asia="宋体", latin="Calibri")
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def setup_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.35

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.3

    for style_name, size, color, before, after in [
        ("Title", 22, RGBColor(22, 50, 79), 0, 14),
        ("Heading 1", 16, RGBColor(46, 116, 181), 18, 10),
        ("Heading 2", 14, RGBColor(46, 116, 181), 14, 8),
        ("Heading 3", 12, RGBColor(31, 77, 120), 10, 6),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.text = "STM32 MP3/WAV 播放器硬件设计学习手册"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=RGBColor(100, 100, 100), east_asia="宋体")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=RGBColor(100, 100, 100), east_asia="宋体")
    add_page_field(footer)
    run = footer.add_run(" 页")
    set_run_font(run, size=9, color=RGBColor(100, 100, 100), east_asia="宋体")


def add_markdown_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, text in enumerate(row):
            set_cell_text(cells[c_idx], text, bold=(r_idx == 0), size=9.5)
            if r_idx == 0:
                set_cell_shading(cells[c_idx], "E8EEF5")
    max_lens = []
    for c in range(len(rows[0])):
        max_lens.append(max(1, max(len(row[c]) for row in rows)))
    widths = column_widths_from_weights(max_lens, total_width_dxa=9360)
    apply_table_geometry(table, widths)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("\n".join(lines))
    set_run_font(run, size=10, east_asia="宋体", latin="Calibri")
    run.font.size = Pt(9.5)


def parse_markdown(doc: Document, source_path: Path):
    lines = source_path.read_text(encoding="utf-8").splitlines()
    i = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(paragraph_buffer).strip()
            if text:
                p = doc.add_paragraph()
                add_inline_runs(p, text)
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            flush_paragraph()
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
        elif line.startswith("!["):
            flush_paragraph()
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                caption, rel = match.groups()
                pic = source_path.parent / rel
                doc.add_picture(str(pic), width=Inches(6.2))
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(caption)
                set_run_font(run, size=9, color=RGBColor(100, 100, 100), east_asia="宋体")
        elif line.startswith("# "):
            flush_paragraph()
            title_text = line[2:].strip()
            if not doc.paragraphs:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(title_text)
            else:
                doc.add_heading(title_text, level=1)
        elif line.startswith("## "):
            flush_paragraph()
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            flush_paragraph()
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("> "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(8)
            add_inline_runs(p, line[2:].strip())
            for run in p.runs:
                set_run_font(run, color=RGBColor(31, 77, 120))
        elif line.startswith("|"):
            flush_paragraph()
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            parsed = []
            for tline in table_lines:
                cells = [cell.strip() for cell in tline.strip("|").split("|")]
                if all(set(cell) <= set("-:") for cell in cells):
                    continue
                parsed.append(cells)
            if parsed:
                add_markdown_table(doc, parsed)
            continue
        elif re.match(r"^\d+\.\s", line):
            flush_paragraph()
            text = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, text)
        elif line.startswith("- [ ] "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, "[ ] " + line[6:])
        elif line.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
        elif not line.strip():
            flush_paragraph()
        else:
            paragraph_buffer.append(line.strip())
        i += 1
    flush_paragraph()


def read_csv(path: Path) -> list[list[str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.reader(f))


def add_landscape_appendix(doc: Document, title: str, rows: list[list[str]], weights: list[float], font_size=7.3):
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    doc.add_heading(title, level=1)
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, text in enumerate(row):
            set_cell_text(cells[c_idx], text, bold=(r_idx == 0), size=font_size)
            if r_idx == 0:
                set_cell_shading(cells[c_idx], "E8EEF5")
    widths = column_widths_from_weights(weights, total_width_dxa=14256)
    apply_table_geometry(table, widths, table_width_dxa=14256)


def build_docx():
    doc = Document()
    setup_document(doc)
    parse_markdown(doc, MANUAL_MD)

    doc.add_page_break()
    add_landscape_appendix(
        doc,
        "附录 A：原理图连接表",
        read_csv(CONNECTION_CSV),
        [0.8, 1.1, 1.0, 0.8, 2.1, 1.5, 2.0],
        font_size=7.0,
    )
    doc.add_page_break()
    add_landscape_appendix(
        doc,
        "附录 B：完整 BOM",
        read_csv(BOM_CSV),
        [1.25, 1.45, 0.45, 0.95, 2.0, 1.8],
        font_size=7.5,
    )

    doc.save(OUTPUT_DOCX)


def write_workbook_from_csv(csv_path: Path, xlsx_path: Path, sheet_name: str, widths: list[float]):
    rows = read_csv(csv_path)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    ws.freeze_panes = "A2"
    ws.sheet_view.showGridLines = False

    header_fill = PatternFill("solid", fgColor="D9EAF7")
    body_fill = PatternFill("solid", fgColor="FFFFFF")
    thin = Side(style="thin", color="C7D3E0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = Font(
                name="Calibri",
                sz=11,
                bold=(row_idx == 1),
                color="000000",
            )
            cell.alignment = Alignment(
                horizontal="center" if row_idx == 1 else "left",
                vertical="center",
                wrap_text=True,
            )
            cell.fill = header_fill if row_idx == 1 else body_fill
            cell.border = border

    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width

    for row_idx, row in enumerate(rows, start=1):
        max_lines = 1
        for col_idx, value in enumerate(row, start=1):
            text = str(value or "")
            width = widths[col_idx - 1]
            estimated_lines = max(1, int(len(text) / max(width * 1.35, 1)) + 1)
            max_lines = max(max_lines, estimated_lines)
        ws.row_dimensions[row_idx].height = 24 if row_idx == 1 else min(18 + (max_lines - 1) * 15, 72)

    ws.auto_filter.ref = ws.dimensions
    wb.save(xlsx_path)


def build_xlsx_tables():
    write_workbook_from_csv(
        CONNECTION_CSV,
        CONNECTION_XLSX,
        "原理图连接表",
        [12, 20, 18, 14, 30, 24, 34],
    )
    write_workbook_from_csv(
        BOM_CSV,
        BOM_XLSX,
        "BOM",
        [24, 26, 8, 16, 34, 30],
    )


def verify_xlsx(path: Path):
    wb = load_workbook(path)
    ws = wb.active
    assert ws.freeze_panes == "A2"
    assert ws.auto_filter.ref == ws.dimensions
    assert all(ws.cell(1, c).font.bold for c in range(1, ws.max_column + 1))
    assert all(ws.cell(2, c).alignment.wrap_text for c in range(1, ws.max_column + 1))


if __name__ == "__main__":
    generate_figures()
    build_docx()
    build_xlsx_tables()
    verify_xlsx(CONNECTION_XLSX)
    verify_xlsx(BOM_XLSX)
    print(OUTPUT_DOCX)
